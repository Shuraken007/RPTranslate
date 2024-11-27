from docx import Document
from pathlib import Path
from ebooklib import epub
from .utils import change_file_name
from .const import CHAPTER_DELIM, REPLACE_AFTER, REPLACE_BEFORE, DOCX

def get_chapter(file):
   arr = []
   chapter_name = None
   for line in file:
      line = line.rstrip()
      if CHAPTER_DELIM in line:
         break
      if len(arr) == 0:
         chapter_name = line
         arr.append('<h1>'+line+'</h1>')   
      else:
         arr.append('<p>'+line+'</p>')

   return '\n'.join(arr), chapter_name

def init_book(s):
   book = epub.EpubBook()
   book.set_title("RuneBoundProfessor")
   book.set_language("ru")
   book.add_author("Acmos")

   #  full_file_name_for_cover = os.path.join(ROOT_DIR,  "bookcovers", book_name + ".jpg")
   #  # send http request and get the image and save it:
   #  urllib.request.urlretrieve(book_cover, full_file_name_for_cover)   
   #  book.set_cover('cover.png', open(full_file_name_for_cover, 'rb').read())
   # add default NCX and Nav file
   book.add_item(epub.EpubNcx())
   book.add_item(epub.EpubNav())
   style = 'BODY {color: white;}'
   nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)   
   book.add_item(nav_css)
   
   return book

def convert_txt_to_epub(s, f_from = REPLACE_AFTER):
   file_path = s.get_out_file(f_from)
   file = open(file_path, 'r')

   book = init_book(s)

   spine = ['nav']
   toc = []   
   number = 0
   while True:
      content, chapter_name = get_chapter(file)
      if not content:
         break
         
      number += 1      
      chapter = epub.EpubHtml(
                        title=chapter_name,
                        file_name='chapter' + str(number) +'.xhtml',
                        lang='ru')
      chapter.set_content('<html><body>' + content + '</body></html>')
      book.add_item(chapter)

      spine.append(chapter)
      toc.append(chapter)

   book.spine = spine
   book.toc = ((epub.Section('Content:'), toc), )

   epub_path = change_file_name(file_path, 'RuneBoundProfessor'+'.epub')
   epub.write_epub(epub_path, book)
   print(epub_path)

   file.close()

def convert_txt_to_docx(s, f_from = REPLACE_BEFORE):
   path = s.get_out_file(f_from)

   document = Document()
   filedata = open(path).read()
   document.add_paragraph(filedata)

   docx_path = Path(path).with_suffix('.docx')
   print(docx_path)
   document.save(docx_path)

def convert_docx_to_txt(s, f_from = DOCX):
   docx_path = s.get_out_file(f_from)
   document = Document(path)
   content = [p.text for p in document.paragraphs]

   txt_path = s.get_out_file(TRANSLATE)
   with open(txt_path, 'w', encoding="utf-8") as f_to:
      f_to.write("\n".join(content))